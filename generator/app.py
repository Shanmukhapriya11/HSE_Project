"""Safety Audit Copilot — evidence-linked, rule-based HSE drafting prototype."""

from __future__ import annotations

import re
from datetime import date, timedelta
from io import BytesIO

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


st.set_page_config(
    page_title="Safety Audit Copilot",
    page_icon=":material/fact_check:",
    layout="wide",
)

DISCLAIMER = (
    "This report is an automatically structured draft based only on the submitted "
    "facility logs. It is not an ISO 45001 certification, regulatory determination, "
    "or final safety assessment. Findings, risk scores, controls and CAPA actions "
    "require verification by qualified HSE personnel."
)

SYNTHETIC_LOG = (
    "During the morning inspection, an unlabeled solvent container was observed in Laboratory A. "
    "No spill or employee exposure was reported. A blocked emergency exit was identified in Storage "
    "Area B and was cleared at 11:20 AM. Two employees reported recurring discomfort while lifting "
    "cartons in the warehouse. The weekly electrical inspection was completed and no visible wiring "
    "defects were identified. Waste containers near the manufacturing area were inspected, but the "
    "log did not state whether segregation labels were present."
)

CATEGORY_RULES = {
    "Chemical": ["solvent", "chemical", "spill", "leak", "acid", "fumes", "container", "label"],
    "Fire and Emergency": ["fire", "smoke", "exit", "evacuation", "extinguisher", "blocked", "emergency"],
    "Electrical": ["wire", "wiring", "electrical", "socket", "voltage", "cable", "inspection"],
    "Ergonomic": ["lifting", "posture", "discomfort", "repetitive", "strain", "workstation"],
    "Mechanical": ["machine", "equipment", "guard", "rotating", "maintenance"],
    "Environmental": ["waste", "segregation", "wastewater", "emissions", "discharge", "recycling"],
    "PPE": ["gloves", "mask", "goggles", "helmet", "protective equipment"],
    "Housekeeping": ["obstruction", "floor", "walkway", "clutter", "trip", "storage"],
}

LOCATIONS = {
    "Laboratory A": ["laboratory a", "laboratory", "lab"],
    "Storage Area B": ["storage area b", "storage area"],
    "Warehouse": ["warehouse"],
    "Manufacturing Area": ["manufacturing area", "manufacturing"],
    "Office": ["office"],
}

DEFAULT_SEVERITY = {
    "Chemical": 4,
    "Fire and Emergency": 5,
    "Electrical": 4,
    "Ergonomic": 3,
    "Mechanical": 4,
    "Environmental": 3,
    "PPE": 3,
    "Housekeeping": 3,
    "Unclassified – Human Review Required": 2,
}

MISSING_QUESTIONS = {
    "Chemical": ["Was the substance identified?", "Was the container isolated?", "Was exposure or a spill confirmed?", "Was the relevant SDS available?"],
    "Fire and Emergency": ["How long was access obstructed?", "Was an evacuation route affected?", "Was the condition formally reported and verified after correction?"],
    "Electrical": ["Was the inspection outcome documented?", "Was isolation required?", "Were defects corrected and verified?"],
    "Ergonomic": ["How many employees were affected?", "Was a manual-handling assessment performed?", "What loads and repetition frequency were involved?"],
    "Mechanical": ["Was the equipment isolated?", "Was guarding verified?", "Was maintenance history reviewed?"],
    "Environmental": ["Was waste correctly labelled and segregated?", "Were quantities recorded?", "Was disposal performed through an approved process?"],
    "PPE": ["Which PPE was required?", "Was it available and correctly used?", "Was its condition checked?"],
    "Housekeeping": ["How long did the condition exist?", "Was access restricted?", "Was correction verified?"],
    "Unclassified – Human Review Required": ["What hazard or control was being assessed?", "What outcome was observed?"],
}

CONTROL_REFERENCES = {
    "Chemical": "CHEM-01 Chemical Labelling",
    "Fire and Emergency": "ERP-02 Emergency Access",
    "Ergonomic": "ERGO-03 Manual Handling",
    "Electrical": "ELEC-04 Electrical Inspection",
    "Environmental": "ENV-05 Waste Segregation",
}


def split_log_into_observations(log_text: str) -> list[str]:
    """Split plain-text logs into non-empty source sentences without rewriting them."""
    normalized = re.sub(r"\s+", " ", log_text.strip())
    if not normalized:
        return []
    return [part.strip() for part in re.split(r"(?<=[.!?])\s+", normalized) if part.strip()]


def classify_observation(sentence: str) -> tuple[str, list[str], str]:
    """Return the category, exact matched keywords, and transparent rule statement."""
    lowered = sentence.lower()
    matches = {
        category: [keyword for keyword in keywords if keyword in lowered]
        for category, keywords in CATEGORY_RULES.items()
    }
    category = max(matches, key=lambda name: len(matches[name]))
    triggered = matches[category]
    if not triggered:
        return "Unclassified – Human Review Required", [], "No configured hazard keyword matched."
    return category, triggered, f"Highest keyword match: {', '.join(triggered)}."


def identify_location(sentence: str) -> tuple[str, str]:
    """Identify only a location explicitly present in the source sentence."""
    lowered = sentence.lower()
    for location, terms in LOCATIONS.items():
        for term in terms:
            if term in lowered:
                return location, f"Location phrase matched: '{term}'."
    return "Not stated", "No configured location phrase matched."


def classify_finding_type(sentence: str) -> tuple[str, str]:
    """Classify wording conservatively; never infer confirmed non-conformity."""
    lowered = sentence.lower()
    unsafe_patterns = ["blocked emergency exit", "unlabeled", "unlabelled", "exposed wiring", "missing machine guard"]
    if any(pattern in lowered for pattern in unsafe_patterns):
        pattern = next(pattern for pattern in unsafe_patterns if pattern in lowered)
        return "Potential Non-Conformity", f"Explicit unsafe-condition phrase matched: '{pattern}'."
    if ("completed" in lowered and ("no visible" in lowered or "no defect" in lowered or "acceptable" in lowered)):
        return "Positive Practice", "Completed check with an explicitly acceptable result."
    if any(term in lowered for term in ["did not state", "not recorded", "not documented", "outcome unknown"]):
        return "Insufficient Evidence", "The source explicitly says a necessary detail or outcome was omitted."
    if any(term in lowered for term in ["recurring", "reported discomfort", "follow-up", "requires review"]):
        return "Follow-Up Required", "The source describes a recurring concern or need for follow-up."
    if any(term in lowered for term in ["observed", "identified", "reported"]):
        return "Safety Observation", "The source explicitly records an observation or report."
    return "Insufficient Evidence", "No explicit condition or completed acceptable outcome was identified."


def risk_level(score: int) -> str:
    if not 1 <= score <= 25:
        raise ValueError("Risk score must be between 1 and 25.")
    if score <= 4:
        return "Low"
    if score <= 9:
        return "Medium"
    if score <= 16:
        return "High"
    return "Critical"


def estimate_risk(category: str, finding_type: str) -> tuple[int, int, int, str, str]:
    """Produce a clearly preliminary category-based estimate, never a verified score."""
    severity = DEFAULT_SEVERITY[category]
    likelihood = {
        "Potential Non-Conformity": 3,
        "Follow-Up Required": 3,
        "Safety Observation": 2,
        "Insufficient Evidence": 2,
        "Positive Practice": 1,
    }[finding_type]
    if finding_type == "Positive Practice":
        severity = 1
    score = likelihood * severity
    rule = f"Preliminary rule-based estimate: likelihood {likelihood} from finding type; severity {severity} from category default."
    return likelihood, severity, score, risk_level(score), rule


def identify_missing_information(sentence: str, category: str, finding_type: str) -> list[str]:
    """Return unanswered questions; never manufacture answers from absent evidence."""
    if finding_type == "Positive Practice":
        return []
    lowered = sentence.lower()
    questions = MISSING_QUESTIONS[category].copy()
    if category == "Chemical" and "no spill" in lowered:
        questions = [q for q in questions if "spill confirmed" not in q.lower()]
    if category == "Ergonomic" and re.search(r"\b(two|2) employees\b", lowered):
        questions = [q for q in questions if "how many employees" not in q.lower()]
    return questions


def recommend_follow_up(category: str, finding_type: str) -> str:
    if finding_type == "Positive Practice":
        return "Retain the inspection record and continue the approved inspection schedule."
    if finding_type == "Insufficient Evidence":
        return f"Obtain and verify the missing {category.lower()} inspection details before concluding."
    if finding_type == "Potential Non-Conformity":
        return f"Have qualified HSE personnel verify the {category.lower()} condition and adequacy of correction."
    return f"Review the {category.lower()} observation with the responsible area and document the outcome."


def generate_capa_suggestion(finding: pd.Series | dict) -> dict[str, str]:
    """Draft actions only; ownership, dates and completion remain deliberately blank."""
    category = finding["Category"]
    finding_id = finding["ID"]
    if finding["Finding Type"] == "Insufficient Evidence":
        immediate = "Collect the missing source information and preserve it with the audit record."
    else:
        immediate = f"Verify the reported {category.lower()} condition and apply a safe interim control if needed."
    return {
        "Finding ID": finding_id,
        "Immediate Correction": immediate,
        "Root-Cause Investigation Area": "Review task, equipment, procedure, training and recent operational changes.",
        "Preventive Action": f"Define and validate a sustainable {category.lower()} control through the site HSE process.",
        "Responsible Person": "",
        "Target Date": "",
        "Verification Status": "Not verified",
    }


def assemble_audit_report(log_text: str) -> pd.DataFrame:
    """Run the modular extraction, classification, evidence and risk pipeline."""
    rows = []
    for index, sentence in enumerate(split_log_into_observations(log_text), start=1):
        category, keywords, category_rule = classify_observation(sentence)
        location, location_rule = identify_location(sentence)
        finding_type, finding_rule = classify_finding_type(sentence)
        likelihood, severity, score, level, risk_rule = estimate_risk(category, finding_type)
        questions = identify_missing_information(sentence, category, finding_type)
        interpretation = f"{finding_type} related to {category.lower()} at {location}."
        rows.append(
            {
                "ID": f"OBS-{index:03d}",
                "Finding Type": finding_type,
                "Category": category,
                "Location": location,
                "Source Evidence": sentence,
                "Interpretation": interpretation,
                "Triggered Rule / Keywords": " | ".join([category_rule, location_rule, finding_rule]),
                "Likelihood": likelihood,
                "Severity": severity,
                "Risk Score": score,
                "Risk Level": level,
                "Risk Basis": risk_rule,
                "Confidence / Extraction Status": "Rule matched" if keywords else "Human review required",
                "Missing Information": "\n".join(f"• {q}" for q in questions) if questions else "No additional question generated by the configured rules.",
                "Suggested Follow-Up": recommend_follow_up(category, finding_type),
                "Internal Checklist / Control Reference": CONTROL_REFERENCES.get(category, "No demo reference assigned"),
                "Status": "Open" if finding_type != "Positive Practice" else "Recorded",
            }
        )
    return pd.DataFrame(rows)


def recalculate_risk(df: pd.DataFrame) -> pd.DataFrame:
    """Validate edits and recalculate all derived risk fields."""
    result = df.copy()
    result["Likelihood"] = pd.to_numeric(result["Likelihood"], errors="coerce").fillna(1).clip(1, 5).astype(int)
    result["Severity"] = pd.to_numeric(result["Severity"], errors="coerce").fillna(1).clip(1, 5).astype(int)
    result["Risk Score"] = result["Likelihood"] * result["Severity"]
    result["Risk Level"] = result["Risk Score"].map(lambda value: risk_level(int(value)))
    result["Risk Basis"] = result.apply(
        lambda row: f"Preliminary estimate using reviewer-editable likelihood {row['Likelihood']} × severity {row['Severity']}.",
        axis=1,
    )
    return result


def risk_cell_style(value: object) -> str:
    """Apply palette-based background colours to risk-level cells."""
    colors = {"Low": "#96CDB0", "Medium": "#EEE8B2", "High": "#C18D52", "Critical": "#5A8F76"}
    return f"background-color: {colors[value]}" if value in colors else ""


def executive_summary(findings: pd.DataFrame) -> str:
    counts = findings["Finding Type"].value_counts()
    high = findings["Risk Level"].isin(["High", "Critical"]).sum()
    return (
        f"The submitted log produced {len(findings)} evidence-linked observations: "
        f"{counts.get('Potential Non-Conformity', 0)} potential non-conformity finding(s), "
        f"{counts.get('Follow-Up Required', 0)} follow-up item(s), and "
        f"{counts.get('Positive Practice', 0)} positive practice(s). "
        f"{high} observation(s) currently have a High/Critical preliminary risk estimate. "
        "All interpretations and estimates require qualified HSE review."
    )


def create_docx(meta: dict, findings: pd.DataFrame, capa: pd.DataFrame) -> bytes:
    """Build a professional in-memory DOCX with evidence and reviewer sign-off."""
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(10)
    title = document.add_heading("Safety Audit Copilot — Audit Draft", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("Synthetic Demo Data", style="Subtitle")
    document.add_heading("Report details", level=1)
    document.add_paragraph(f"Site: {meta['site']}\nAudit date: {meta['date']}\nScope: {meta['scope']}\nAuditor/reviewer: {meta['reviewer'] or 'Not assigned'}")
    document.add_heading("Executive summary", level=1)
    document.add_paragraph(executive_summary(findings))
    document.add_heading("Findings summary", level=1)
    summary = document.add_table(rows=1, cols=5)
    for cell, text in zip(summary.rows[0].cells, ["ID", "Type", "Category", "Risk", "Status"]):
        cell.text = text
    for _, row in findings.iterrows():
        cells = summary.add_row().cells
        for cell, value in zip(cells, [row["ID"], row["Finding Type"], row["Category"], f"{row['Risk Score']} — {row['Risk Level']}", row["Status"]]):
            cell.text = str(value)
    document.add_heading("Detailed findings and source evidence", level=1)
    for _, row in findings.iterrows():
        document.add_heading(f"{row['ID']} — {row['Finding Type']}", level=2)
        document.add_paragraph(f"Source evidence: {row['Source Evidence']}")
        document.add_paragraph(f"Interpretation: {row['Interpretation']}")
        document.add_paragraph(f"Evidence rule: {row['Triggered Rule / Keywords']}")
        document.add_paragraph(f"Preliminary risk: {row['Likelihood']} × {row['Severity']} = {row['Risk Score']} ({row['Risk Level']})")
        document.add_paragraph(f"Missing-information questions: {row['Missing Information']}")
        document.add_paragraph(f"Suggested follow-up: {row['Suggested Follow-Up']}")
        document.add_paragraph(f"Internal checklist / control reference: {row['Internal Checklist / Control Reference']}")
    document.add_heading("Draft CAPA tracker", level=1)
    if capa.empty:
        document.add_paragraph("No draft CAPA items were generated.")
    else:
        table = document.add_table(rows=1, cols=len(capa.columns))
        for cell, column in zip(table.rows[0].cells, capa.columns):
            cell.text = column
        for _, row in capa.iterrows():
            for cell, value in zip(table.add_row().cells, row):
                cell.text = str(value)
    document.add_heading("Limitations", level=1)
    document.add_paragraph(DISCLAIMER)
    document.add_heading("Human reviewer approval", level=1)
    document.add_paragraph(
        "Reviewer name: __________________________________________\n\n"
        "Decision/comments: _______________________________________\n"
        "__________________________________________________________\n\n"
        "Signature: ________________________    Date: _______________"
    )
    footer = section.footer.paragraphs[0]
    footer.text = "Automatically structured draft — qualified HSE verification required."
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def capa_from_findings(findings: pd.DataFrame) -> pd.DataFrame:
    relevant = findings[findings["Finding Type"] != "Positive Practice"]
    columns = ["Finding ID", "Immediate Correction", "Root-Cause Investigation Area", "Preventive Action", "Responsible Person", "Target Date", "Verification Status"]
    return pd.DataFrame([generate_capa_suggestion(row) for _, row in relevant.iterrows()], columns=columns)


def initialize_state() -> None:
    st.session_state.setdefault("facility_log", "")
    st.session_state.setdefault("findings", None)
    st.session_state.setdefault("generated_meta", None)


def load_synthetic_example() -> None:
    """Populate the log before Streamlit instantiates its keyed widget."""
    st.session_state.facility_log = SYNTHETIC_LOG
    st.session_state.findings = None
    st.session_state.generated_meta = None


initialize_state()

st.title("Safety Audit Copilot")
st.caption("Evidence-linked HSE audit drafting prototype · Synthetic Demo Data")
st.warning(DISCLAIMER, icon=":material/warning:")

with st.container(border=True):
    st.markdown("**How it works**")
    st.caption("Logs → extraction → classification → evidence linking → preliminary prioritization → audit draft → HSE review")

# Keep the workflow on one continuous page without navigation tabs.
input_tab, findings_tab, capa_tab, report_tab = (
    st.container(),
    st.container(),
    st.container(),
    st.container(),
)

with input_tab:
    st.subheader("Facility log input")
    top_left, top_right = st.columns(2)
    site = top_left.text_input("Facility/site name")
    audit_date = top_right.date_input("Audit date", value=date.today())
    scope = st.text_input("Audit scope")
    reviewer = st.text_input("Auditor/reviewer name")
    log_text = st.text_area("Unstructured daily facility logs", key="facility_log", height=220, placeholder="Paste facility logs here. No evidence will be invented beyond this text.")
    with st.container(horizontal=True):
        st.button(
            "Load synthetic example",
            icon=":material/science:",
            on_click=load_synthetic_example,
        )
        generate = st.button("Generate Audit Draft", type="primary", icon=":material/description:")
    if generate:
        if not log_text.strip():
            st.error("Enter facility logs or load the synthetic example before generating a draft.")
        else:
            generated = assemble_audit_report(log_text)
            if generated.empty:
                st.error("No observations could be extracted from the submitted text.")
            else:
                st.session_state.findings = generated
                st.session_state.generated_meta = {"site": site.strip() or "Not stated", "date": audit_date.isoformat(), "scope": scope.strip() or "Not stated", "reviewer": reviewer.strip()}
                st.success(f"Generated {len(generated)} evidence-linked observations. The draft appears below.")

findings = st.session_state.findings

with findings_tab:
    st.subheader("Extracted findings")
    if findings is None:
        st.info("Generate an audit draft from the facility log above first.")
    else:
        st.caption("Likelihood and severity are editable from 1–5. Status is also editable. Risk score and level are recalculated from those inputs and remain preliminary.")
        display_columns = ["ID", "Finding Type", "Category", "Location", "Source Evidence", "Likelihood", "Severity", "Risk Score", "Risk Level", "Status"]
        edited = st.data_editor(
            findings[display_columns],
            key="findings_editor",
            hide_index=True,
            width="stretch",
            disabled=["ID", "Finding Type", "Category", "Location", "Source Evidence", "Risk Score", "Risk Level"],
            column_config={
                "Likelihood": st.column_config.NumberColumn(min_value=1, max_value=5, step=1),
                "Severity": st.column_config.NumberColumn(min_value=1, max_value=5, step=1),
                "Risk Score": st.column_config.NumberColumn(format="%d"),
                "Status": st.column_config.SelectboxColumn(options=["Open", "In review", "Resolved", "Recorded"]),
            },
        )
        updated = findings.copy()
        updated.update(edited)
        updated = recalculate_risk(updated)
        st.session_state.findings = updated
        findings = updated
        st.caption("Calculated preview after edits")
        st.dataframe(
            findings[["ID", "Likelihood", "Severity", "Risk Score", "Risk Level", "Status"]].style.map(risk_cell_style),
            hide_index=True,
            width="stretch",
        )
        st.caption("All automated scores are preliminary rule-based estimates, not verified risk assessments.")
        for _, row in findings.iterrows():
            with st.expander(f"{row['ID']} · Evidence and reasoning", icon=":material/link:"):
                st.markdown(f"**Exact source sentence:** {row['Source Evidence']}")
                st.markdown(f"**Interpretation:** {row['Interpretation']}")
                st.markdown(f"**Triggered rule / keywords:** {row['Triggered Rule / Keywords']}")
                st.markdown(f"**Extraction status:** {row['Confidence / Extraction Status']}")
                st.markdown(f"**Risk basis:** {row['Risk Basis']}")
                st.markdown(f"**Suggested follow-up:** {row['Suggested Follow-Up']}")
                st.markdown(f"**Internal checklist / control reference:** {row['Internal Checklist / Control Reference']}")

with capa_tab:
    st.subheader("CAPA and information gaps")
    if findings is None:
        st.info("Generate an audit draft first.")
    else:
        for _, row in findings.iterrows():
            with st.container(border=True):
                st.markdown(f"**{row['ID']} · {row['Category']}**")
                st.markdown(row["Missing Information"])
                st.caption("Questions are intentionally unanswered unless the source log supplies evidence.")
        capa = capa_from_findings(findings)
        st.subheader("Draft CAPA tracker")
        st.caption("Responsible person and target date are deliberately blank for human assignment.")
        st.dataframe(capa, hide_index=True, width="stretch")

with report_tab:
    st.subheader("Audit report preview")
    if findings is None or st.session_state.generated_meta is None:
        st.info("Generate an audit draft first.")
    else:
        meta = st.session_state.generated_meta
        capa = capa_from_findings(findings)
        with st.container(border=True):
            st.header("Safety Audit Copilot — Audit Draft")
            st.caption("Synthetic Demo Data")
            st.markdown(f"**Site:** {meta['site']}  \n**Date:** {meta['date']}  \n**Scope:** {meta['scope']}  \n**Auditor/reviewer:** {meta['reviewer'] or 'Not assigned'}")
            st.subheader("Executive summary")
            st.write(executive_summary(findings))
            st.subheader("Findings summary")
            with st.container(horizontal=True):
                st.metric("Total observations", len(findings), border=True)
                st.metric("Potential non-conformities", int((findings["Finding Type"] == "Potential Non-Conformity").sum()), border=True)
                st.metric("Follow-ups required", int((findings["Finding Type"] == "Follow-Up Required").sum()), border=True)
                st.metric("High/Critical preliminary risks", int(findings["Risk Level"].isin(["High", "Critical"]).sum()), border=True)
                st.metric("Unresolved information gaps", int((findings["Missing Information"] != "No additional question generated by the configured rules.").sum()), border=True)
                st.metric("Positive practices", int((findings["Finding Type"] == "Positive Practice").sum()), border=True)
            st.subheader("Detailed findings, evidence and preliminary risk")
            for _, row in findings.iterrows():
                st.markdown(f"**{row['ID']} — {row['Finding Type']}**")
                st.write(row["Interpretation"])
                st.caption(f"Source evidence: {row['Source Evidence']}")
                st.caption(f"Preliminary risk: {row['Likelihood']} × {row['Severity']} = {row['Risk Score']} ({row['Risk Level']})")
                st.markdown(f"Missing information:  \n{row['Missing Information']}")
            st.subheader("Draft CAPA tracker")
            st.dataframe(capa, hide_index=True, width="stretch")
            st.subheader("Limitations")
            st.warning(DISCLAIMER, icon=":material/warning:")
            st.subheader("Human reviewer approval")
            st.text(
                "Reviewer name: __________________________________________\n\n"
                "Decision/comments: _______________________________________\n"
                "__________________________________________________________\n\n"
                "Signature: ________________________    Date: _______________"
            )
        st.info("A production implementation could map findings to an organization’s licensed ISO 45001 requirements, internal procedures and legal register after expert validation.")
        docx_bytes = create_docx(meta, findings, capa)
        csv_bytes = findings[["ID", "Finding Type", "Category", "Location", "Source Evidence", "Likelihood", "Severity", "Risk Score", "Risk Level", "Status"]].to_csv(index=False).encode("utf-8")
        with st.container(horizontal=True):
            st.download_button("Download DOCX", docx_bytes, file_name="safety_audit_draft.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", icon=":material/download:")
            st.download_button("Download findings CSV", csv_bytes, file_name="safety_audit_findings.csv", mime="text/csv", icon=":material/download:")
